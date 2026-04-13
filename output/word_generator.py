"""
output/word_generator.py — Generador de Informes Profesionales en Word (.docx)
"""

import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generar_informe_word(
    cuit: str,
    periodo: str,
    provincia_id: str,
    volumen_ventas_anual: float,
    resultados_calc: list,
    justificacion_llm: str,
    situacion_especial: str | None = None,
    output_path: str = "informe.docx"
):
    """
    Crea un documento .docx profesional con el dictamen de auditoría.
    """
    doc = Document()

    # 1. Encabezado / Título
    title = doc.add_heading('INFORME TÉCNICO DE AUDITORÍA FISCAL', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. Información del Cliente
    doc.add_heading('1. Información General', level=1)
    p = doc.add_paragraph()
    p.add_run(f"CUIT del Contribuyente: ").bold = True
    p.add_run(f"{cuit}\n")
    p.add_run(f"Período Analizado: ").bold = True
    p.add_run(f"{periodo}\n")
    p.add_run(f"Jurisdicción: ").bold = True
    p.add_run(f"{provincia_id.upper()}\n")
    p.add_run(f"Volumen de Ventas Anual: ").bold = True
    p.add_run(f"${volumen_ventas_anual:,.2f}")

    # 3. Situación Especial (Si existe)
    if situacion_especial:
        doc.add_heading('2. Contexto Específico', level=1)
        doc.add_paragraph(situacion_especial)

    # 4. Detalle de Actividades y Alícuotas
    doc.add_heading('3. Determinación de Alícuotas (Ley Impositiva Anual)', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'NAES'
    hdr_cells[1].text = 'Actividad / Descripción'
    hdr_cells[2].text = 'Tasa Base'
    hdr_cells[3].text = 'Tasa Sugerida'

    for r in resultados_calc:
        row_cells = table.add_row().cells
        row_cells[0].text = str(r.naes_encontrado)
        row_cells[1].text = str(r.actividad_desc_norma)
        row_cells[2].text = f"{r.alicuota_base}%"
        row_cells[3].text = f"{r.alicuota_final}%"

    # 5. Dictamen del Auditor (Justificación LLM)
    doc.add_heading('4. Dictamen y Fundamentación Legal', level=1)

    # Limpiar la justificación: si es un error técnico, mostrar solo mensaje limpio
    ERROR_KEYWORDS = ['RESOURCE_EXHAUSTED', 'NOT_FOUND', 'INVALID_ARGUMENT',
                      'quotaMetric', 'type.googleapis', 'Traceback', 'Exception']
    is_raw_error = any(kw in justificacion_llm for kw in ERROR_KEYWORDS)
    
    if is_raw_error:
        # Extraer solo la primera línea significativa del error
        first_line = justificacion_llm.split('\n')[0][:300]
        doc.add_paragraph(
            f"⚠️ NOTA: El Motor de IA no pudo generar el dictamen en este procesamiento.\n"
            f"Motivo: {first_line}\n\n"
            f"Por favor, reprocese este contribuyente o contacte al administrador del sistema."
        )
    else:
        # Texto normal — dividir en párrafos para que quede bien en Word
        paragraphs = justificacion_llm.split('\n')
        for para in paragraphs:
            if para.strip():
                if para.strip().startswith('###'):
                    clean_text = para.replace('###', '').strip()
                    p = doc.add_paragraph()
                    p.add_run(clean_text).bold = True
                else:
                    doc.add_paragraph(para)

    # 6. Pie de página y Firma
    doc.add_paragraph("\n---")
    
    # Firma Institucional LL
    sig = doc.add_paragraph()
    sig.add_run("Auditado por:\n").bold = True
    sig.add_run("Auditor Fiscal Agente de LL (Lisicki Litvin)\n").bold = True
    sig.add_run("Departamento de Tax Compliance / Auditoría Externa").italic = True
    sig.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("\n")
    footer = doc.add_paragraph(f"Documento generado de forma automática por Agente Fiscal Senior LL - {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Guardar
    doc.save(output_path)
    return output_path

